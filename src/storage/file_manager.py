from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Union

import numpy as np
import soundfile as sf


# ── helpers ──────────────────────────────────────────────────────────────────

def ensure_directory(path: Union[str, Path]) -> Path:
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


def get_save_path(extension: str, folder: Union[str, Path] = None) -> Path:
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    directory = Path(folder) if folder else Path.home() / "Recordings"
    ensure_directory(directory)
    return directory / f"{timestamp}.{extension.lstrip('.')}"


# ── raw save (used internally) ────────────────────────────────────────────────

def export_wav(audio_data: np.ndarray, sample_rate: int, filepath: Union[str, Path]) -> Path:
    """Write in-memory audio to a WAV file (PCM 16-bit)."""
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    sf.write(str(filepath), audio_data, sample_rate, subtype='PCM_16')
    return filepath


def save_wav(audio_data: np.ndarray, sample_rate: int, filepath: Union[str, Path]) -> Path:
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    sf.write(str(filepath), audio_data, sample_rate)
    return filepath


def save_txt(text: str, filepath: Union[str, Path]) -> Path:
    """Save raw text string to file (used by Ctrl+S quick-save)."""
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    filepath.write_text(text, encoding="utf-8")
    return filepath


# ── export functions (take segment list) ─────────────────────────────────────
# Segment format: {"timestamp": "00:01:23", "text": str, "translation": str}

def export_txt(segments: List[dict], filepath: Union[str, Path]) -> Path:
    """Export segments as timestamped plain text (combined transcript + translation)."""
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    lines = []
    for seg in segments:
        lines.append(f"[{seg['timestamp']}] {seg['text']}\n")
        if seg.get("translation"):
            lines.append(f"[{seg['timestamp']}] {seg['translation']}\n")
        lines.append("\n")
    filepath.write_text("".join(lines), encoding="utf-8")
    return filepath


def export_transcript_txt(segments: List[dict], filepath: Union[str, Path]) -> Path:
    """Export transcript only as timestamped plain text."""
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    lines = []
    for seg in segments:
        lines.append(f"[{seg['timestamp']}] {seg['text']}\n")
    filepath.write_text("".join(lines), encoding="utf-8")
    return filepath


def export_translation_txt(segments: List[dict], filepath: Union[str, Path]) -> Path:
    """Export translation only as timestamped plain text."""
    filepath = Path(filepath)
    ensure_directory(filepath.parent)
    lines = []
    for seg in segments:
        if seg.get("translation"):
            lines.append(f"[{seg['timestamp']}] {seg['translation']}\n")
    filepath.write_text("".join(lines), encoding="utf-8")
    return filepath


def export_docx(
    segments: List[dict],
    filepath: Union[str, Path],
    title: str = "Recording Transcript",
) -> Path:
    """Export segments as a formatted Word document."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        tcPr.append(shd)

    filepath = Path(filepath)
    ensure_directory(filepath.parent)

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.runs[0].font.color.rgb = RGBColor(0x21, 0x96, 0xF3)

    doc.add_paragraph(f"Recording date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("Time", "Transcript", "Translation")):
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "2196F3")

    # Set column widths (header row)
    for cell, width in zip(hdr, (Cm(1.5), Cm(8), Cm(8))):
        cell.width = width

    # Data rows
    for i, seg in enumerate(segments):
        row = table.add_row()
        row.cells[0].text = seg["timestamp"]
        row.cells[1].text = seg["text"]
        row.cells[2].text = seg.get("translation", "")
        for cell, width in zip(row.cells, (Cm(1.5), Cm(8), Cm(8))):
            cell.width = width
        if i % 2 == 1:
            for cell in row.cells:
                _set_cell_bg(cell, "F5F5F5")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by RealTime Recorder")
    if footer.runs:
        footer.runs[0].italic = True

    doc.save(str(filepath))
    return filepath


def export_srt(segments: List[dict], filepath: Union[str, Path]) -> Path:
    """Export segments as SRT subtitle file."""

    def _srt_time(total_secs: int) -> str:
        h = total_secs // 3600
        m = (total_secs % 3600) // 60
        s = total_secs % 60
        return f"{h:02d}:{m:02d}:{s:02d},000"

    def _ts_to_secs(ts: str) -> int:
        parts = ts.split(":")
        return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])

    filepath = Path(filepath)
    ensure_directory(filepath.parent)

    lines = []
    for i, seg in enumerate(segments):
        start = _ts_to_secs(seg["timestamp"])
        if i + 1 < len(segments):
            end = _ts_to_secs(segments[i + 1]["timestamp"])
        else:
            end = start + 4

        lines.append(str(i + 1))
        lines.append(f"{_srt_time(start)} --> {_srt_time(end)}")
        lines.append(seg["text"])
        if seg.get("translation"):
            lines.append(f"[{seg['translation']}]")
        lines.append("")

    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath


def export_mp3(wav_filepath: Union[str, Path], mp3_filepath: Union[str, Path]) -> Path:
    """Convert WAV file to MP3 at 128k using pydub + ffmpeg."""
    try:
        from pydub import AudioSegment
    except ImportError:
        raise RuntimeError("pydub is not installed")

    try:
        audio = AudioSegment.from_wav(str(wav_filepath))
        mp3_filepath = Path(mp3_filepath)
        ensure_directory(mp3_filepath.parent)
        audio.export(str(mp3_filepath), format="mp3", bitrate="128k")
        return mp3_filepath
    except Exception as e:
        msg = str(e).lower()
        if "ffmpeg" in msg or "avconv" in msg or "couldn't find" in msg:
            raise RuntimeError("Install ffmpeg to export MP3: brew install ffmpeg")
        raise RuntimeError(f"MP3 export failed: {e}")


def export_mp3_from_data(
    audio_data: np.ndarray,
    sample_rate: int,
    mp3_filepath: Union[str, Path],
) -> Path:
    """Convert in-memory audio to MP3 at 128k using pydub + ffmpeg."""
    try:
        from pydub import AudioSegment
    except ImportError:
        raise RuntimeError("pydub is not installed")

    try:
        # pydub works with int16 PCM bytes
        pcm = (audio_data * 32767).astype(np.int16)
        audio = AudioSegment(
            pcm.tobytes(),
            frame_rate=sample_rate,
            sample_width=2,  # 16-bit = 2 bytes
            channels=1,
        )
        mp3_filepath = Path(mp3_filepath)
        ensure_directory(mp3_filepath.parent)
        audio.export(str(mp3_filepath), format="mp3", bitrate="128k")
        return mp3_filepath
    except Exception as e:
        msg = str(e).lower()
        if "ffmpeg" in msg or "avconv" in msg or "couldn't find" in msg:
            raise RuntimeError("Install ffmpeg to export MP3: brew install ffmpeg")
        raise RuntimeError(f"MP3 export failed: {e}")
